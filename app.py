from flask import Flask, render_template, request, send_file, jsonify
from openai import OpenAI
import re
import markdown
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor
import json
import os
from io import BytesIO
import logging
import sys

from googleapiclient.discovery import build

import time

logging.basicConfig(level=logging.DEBUG, format="%(asctime)s - %(levelname)s - %(message)s", handlers=[
    logging.StreamHandler(sys.stdout),
    logging.StreamHandler(sys.stderr)
])

open_ai_api = os.getenv("OPEN_AI_API")
if open_ai_api != None:
    client = OpenAI(api_key = open_ai_api)
else:
    logging.info("OPEN AI API not found")

logging.info("open_ai_api")


app = Flask(__name__)

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/submit', methods=['POST'])
def submit():
    topic_name = request.form['topic_name']
    region = request.form['region']

    query = topic_name + " " + region + " Latest NEWS"

    source_links = google_results(query)

    first_chat_instructions = "You are an intelligent assistant tasked with analyzing internet activity related to a specific target (person, business, or organization) over the last 7 days. Your goal is to provide a high-level summary and structured output that highlights the most interesting and relevant insights, which will serve as a foundation for creating detailed individual reports."

    # first_prompt = f'You are a highly skilled internet analyst tasked with analyzing the online presence and activity of a specific target over the last 7 days. Your job is to produce a detailed, engaging, and structured analysis that provides actionable insights, measurable comparisons, and psychometric trends. This analysis will form the foundation for a comprehensive multi-page report. Target Details: Name: {topic_name}. Region: {region}. Timeframe: Analyze data specifically from the last 7 days. Instructions: Generate 9-11 Unique Topics: Identify at least 9-11 engaging and detailed topics related to the target’s online activity during the specified timeframe. Ensure each topic reflects key developments, public reactions, or noteworthy trends involving the target. The headings must be creative, specific, and newsworthy (e.g., “How [Target Name] Is Transforming Public Perception About [Topic]” or “The Surprising Impact of [Target Name]\'s Recent Actions on [Region]”). Avoid generic or repetitive titles—focus on highlighting unique insights and trends. Content Requirements for Each Topic: Detailed Summary: Provide an in-depth overview of the topic, including relevant activities or developments. Numerical Data: Include specific, measurable metrics such as: Engagement growth rates (e.g., “Engagement increased by 25% compared to last week”). Sentiment distribution percentages (e.g., “40% positive, 35% negative, and 25% neutral”). Platform-specific activity (e.g., “60% of mentions on Twitter, 25% on Facebook, and 15% on Instagram”). Comparative Insights: Highlight how the target’s performance compares to: Previous weeks. Competitors or peers in the same region or field. Broader trends within the industry or region. Psychometric Analysis: Analyze public sentiment and emotional responses: Identify dominant emotions like trust, anger, joy, or fear. Highlight any shifts in sentiment or behavioral trends. Tangible Impacts: Assess the measurable effects of the target’s activities: Changes in public opinion, reputation, or brand perception. Business outcomes, policy shifts, or other significant impacts. Highlight Recurring Themes and Trends: Identify major keywords, hashtags, or recurring phrases associated with the target. Specify which platforms (e.g., Twitter, Facebook, Instagram) saw the most activity for each theme. Predictions for the Future: Based on current data, provide forecasts on how public sentiment, activity levels, or the target’s influence might evolve in the coming weeks. Include data-driven trends and potential scenarios. Special Instructions: Ensure each topic has enough depth and detail to support a multi-page report. The output should include insightful, numerical, and comparative analysis where applicable. Avoid generic or overly broad headings. Focus on creating specific, impactful, and engaging topics tailored to the target. Clearly indicate if certain data points are unavailable, but maintain a comprehensive and analytical approach. Use a professional tone suitable for an executive-level audience. Output Format: Structure the response into a JSON object. For each heading, include subheadings and a detailed summary incorporating all specified data points.Output Format: Provide the output in the following structured format: {{ "target": "[Insert Name]", "entity_type": "[Person/Business/Organization]", "timeframe": "Last 7 days", "content_overview": [ {{ "heading": "[Heading Title 1]", "subheadings": [ "[Subheading 1.1]", "[Subheading 1.2]" ], "summary": "[Brief summary of this section and why it is relevant.]" }}, {{ "heading": "[Heading Title 2]", "subheadings": [ "[Subheading 2.1]", "[Subheading 2.2]" ], "summary": "[Brief summary of this section and why it is relevant.]" }} ] }}'

    first_part_prompt = f""""You are a highly skilled internet analyst tasked with analyzing the online presence and activity of a specific target over the last 7 days. Your job is to produce a detailed, engaging, and structured analysis that provides actionable insights, measurable comparisons, and psychometric trends. This analysis will form the foundation for a comprehensive multi-page report. Source Limitation: Use only the information explicitly available in the source links provided below. Do not reference or incorporate any external knowledge, prior understanding, or information from other sources. All insights must be derived solely from the content of the provided links. Source Links: """
    second_half_prompt = f"""Target Details: Name: {topic_name}. Region: {region}. Instructions: Generate 9-11 Unique Topics: Identify at least 9-11 engaging and detailed topics related to the target’s online activity during the specified timeframe. Ensure each topic reflects key developments, public reactions, or noteworthy trends involving the target. The headings must be creative, specific, and newsworthy (e.g., “How [Target Name] Is Transforming Public Perception About [Topic]” or “The Surprising Impact of [Target Name]’s Recent Actions on [Region]”). Avoid generic or repetitive titles—focus on highlighting unique insights and trends. Content Requirements for Each Topic: Detailed Summary: Provide an in-depth overview of the topic, including relevant activities or developments strictly sourced from the links. Numerical Data: Include specific, measurable metrics such as: Engagement growth rates (e.g., “Engagement increased by 25% compared to last week”). Sentiment distribution percentages (e.g., “40% positive, 35% negative, and 25% neutral”). Platform-specific activity (e.g., “60% of mentions on Twitter, 25% on Facebook, and 15% on Instagram”). Comparative Insights: Highlight how the target’s performance compares to: Previous weeks (if this data is available in the provided sources). Competitors or peers in the same region or field. Broader trends within the industry or region. Psychometric Analysis: Analyze public sentiment and emotional responses: Identify dominant emotions like trust, anger, joy, or fear. Highlight any shifts in sentiment or behavioral trends. Tangible Impacts: Assess the measurable effects of the target’s activities: Changes in public opinion, reputation, or brand perception. Business outcomes, policy shifts, or other significant impacts. Highlight Recurring Themes and Trends: Identify major keywords, hashtags, or recurring phrases associated with the target. Specify which platforms (e.g., Twitter, Facebook, Instagram) saw the most activity for each theme. Predictions for the Future: Based on current data, provide forecasts on how public sentiment, activity levels, or the target’s influence might evolve in the coming weeks. Include data-driven trends and potential scenarios. Special Instructions: Ensure each topic has enough depth and detail to support a multi-page report. The output should include insightful, numerical, and comparative analysis where applicable. Clearly indicate if certain data points are unavailable, but maintain a comprehensive and analytical approach. Use a professional tone suitable for an executive-level audience.Output Format: Structure the response into a JSON object. For each heading, include subheadings and a detailed summary incorporating all specified data points. json Copy code {{ "target": "[Insert Name]", "entity_type": "[Person/Business/Organization]", "timeframe": "Last 7 days", "content_overview": [ {{ "heading": "[Heading Title 1]", "subheadings": [ "[Subheading 1.1]", "[Subheading 1.2]" ], "summary": "[Brief summary of this section and why it is relevant.]" }}, {{ "heading": "[Heading Title 2]", "subheadings": [ "[Subheading 2.1]", "[Subheading 2.2]" ], "summary": "[Brief summary of this section and why it is relevant.]" }} ] }} Reminder: Adhere strictly to the information in the provided links. Any data or insight not explicitly found within the links must be omitted."""

    first_page_prompt = f"You are an advanced data analyst tasked with creating an engaging and insightful introduction for the first page of a comprehensive report. Your goal is to analyze the content from the provided links and produce a captivating, data-rich narrative that highlights key insights, trends, and measurable impacts. The aim is to immediately grab the reader's attention and set the tone for the rest of the document. Topic should be {topic_name} and related to {region}. Source Links:"
    
    first_page_prompt_2 = f"""Instructions: Analyze the Links: Extract key trends, data points, and recurring themes across all the sources. Focus on areas where measurable impacts are evident, such as increases or decreases in engagement, sentiment, or influence. Present Analytical Insights: Provide a mix of qualitative insights and quantitative data to make the introduction compelling. Use percentages, growth rates, or probabilities wherever applicable (e.g., "Engagement increased by 25% in the last week," "Public sentiment shifted by 40% in favor of X"). Include Impactful Highlights: Summarize how the person, organization, or topic influenced the region or audience within the timeframe. Highlight tangible outcomes, such as changes in public perception, policy developments, or increased awareness around specific themes. Make It Interesting to Read: Use dynamic and vivid language to make the content engaging and relatable. Avoid overly technical or dense phrasing; instead, prioritize clear, compelling storytelling supported by data. Identify Key Themes and Trends: Summarize the major topics and themes that emerge from the sources, such as public sentiment, engagement growth, or focus on specific issues. Briefly describe why these themes are relevant or important. Keep It Compact and Focused: Ensure the content is concise yet informative, suitable for a professional audience. Example Content Flow: Start with a compelling hook or headline summarizing the overall findings. Follow with 2-3 sentences describing the primary insights, supported by numerical data. Highlight one or two significant impacts or trends that stand out from the analysis. Conclude with a forward-looking statement or key takeaway that encourages the reader to explore further. Reminder: Use only the provided links as your source of information. Ensure the first page content is engaging, analytical, and provides enough detail to captivate the reader and encourage them to delve deeper into the document."""

    first_page_prompt_final = first_page_prompt + str(source_links) + first_page_prompt_2
    first_page_response = chat_assist(first_chat_instructions, first_page_prompt_final)

    first_prompt = first_part_prompt + str(source_links) + second_half_prompt

    response = chat_assist(first_chat_instructions, first_prompt)

    detail_adder_instructions = "Understand the Context: Recognize that the goal is to expand a brief summary, heading, and subheadings into a fully developed section for a report. Each section must provide sufficient depth and insights to stand as a standalone page in a multi-page document. Follow a Logical Structure: For each section: Begin with an introduction that provides context based on the heading and summary. Develop each subheading into a detailed subsection with insights, data, and analysis. Conclude the section with a summary of key takeaways and any broader implications or predictions. Incorporate Data and Analysis: Where applicable: Use quantitative insights from the summary, such as engagement rates, sentiment distribution, or platform metrics. Provide comparative analysis (e.g., with competitors, past trends, or regional benchmarks). Highlight psychometric insights, such as emotional responses or behavioral patterns. Maintain Relevance and Depth: Ensure the content stays relevant to the heading and summary. Expand on the subheadings to add depth, avoiding repetition or generic descriptions. Write Professionally: Use a formal, executive-level tone suitable for a report. Ensure the content flows logically between subsections. Address Missing Data: If specific data points (e.g., numerical metrics or comparative insights) are not available, explicitly state it (e.g., “Detailed data on competitor engagement is unavailable for this period”). Content Length: Generate approximately 300-500 words per section, ensuring enough detail to create a full page of content while maintaining clarity and focus. Be Engaging and Informative: Ensure the content is not only informative but also engaging. Use storytelling elements or highlight the significance of trends to make the analysis compelling."

    # graph_creator_instructions = 'You are a professional report generator tasked with analyzing detailed data and creating appropriate visual elements for each section of a report. You will be provided with a heading, subheadings, and summary. Your task is to determine whether to generate graphs (based on numerical data) or retrieve stock images (based on context). Input Data: Heading: [Insert heading text] Subheadings: [Insert list of subheadings] Summary: [Insert detailed summary with numerical data and context] Instructions: Determine Visualization Type: If the input contains numerical data, suggest one or more graphs that can visually represent the data. If no numerical data is available, identify relevant stock images based on keywords extracted from the heading and summary. If both graphs and images are relevant, include both. Graph Creation: Specify the type of graph (e.g., pie chart, bar chart, line graph) and what it represents. Provide details for the graph, including labels, titles, and numerical data points. Stock Image Retrieval: Suggest 2-3 keywords based on the heading and summary for searching professional images. If possible, include URLs for related stock images from Unsplash or similar sources. Output Format: Maintain consistent JSON formatting with the following keys: heading: The section heading. visualization_type: Indicate whether the visualization is “graph”, “stock_images”, or “both”. graphs: If applicable, include: graph_type: The type of graph (e.g., pie, bar). description: A brief explanation of the graph’s purpose. data: Numerical data used to generate the graph. images: If applicable, include: keywords: Keywords for image search. urls: URLs of suggested stock images. Special Instructions: If generating graphs, ensure the data is presented clearly and accurately. If suggesting images, ensure the keywords align with the context of the heading. For sections with no suitable visuals, clearly state that no graphs or images are applicable.'

    # graph_prompt = 'You are a professional report generator tasked with analyzing detailed data and creating appropriate visual elements for each section of a report. Your job is to determine whether to generate graphs (based on numerical data) or suggest stock images (based on context). Your goal is to ensure that each visual enhances the report’s clarity and presentation. Input Data: Heading: [Insert heading text] Subheadings: [Insert list of subheadings] Summary: [Insert detailed summary with numerical data and context] Instructions: Determine Visualization Type: If numerical data is present or can be inferred, suggest graphs (e.g., pie charts, bar charts, line graphs). If no numerical data is available, suggest relevant stock images based on keywords extracted from the heading and summary. If both graphs and images are applicable, include both. Graph Suggestions: Specify the type of graph and its purpose. Include labels, titles, and numerical data points. Stock Image Suggestions: Provide 2-3 keywords based on the heading and summary for searching professional images. Include URLs for related stock images if possible. Output Format: heading: The section heading. visualization_type: Indicate “graph,” “stock_images,” or “both.” graphs: If applicable, include: graph_type: The type of graph. description: Explanation of its purpose. data: Numerical data points. images: If applicable, include: keywords: Keywords for image search. urls: Suggested stock image URLs. Special Instructions: Ensure the visuals are directly relevant to the content and enhance clarity. If neither graphs nor images are suitable, state: “No suitable visual content available.” Maintain professional and consistent formatting for integration into the report.'

    pattern = r"^.*?```json| ```\s.*$"
    str_json_response = re.sub( pattern, "", response)
    str_json_response = str_json_response.replace("```", "")
    try:
        json_response = json.loads(str_json_response)
    except json.JSONDecodeError as e:

        return f"Error parsing JSON: {e}\n This is the str JSON: {str_json_response}"
    try:
        full_report = "" + first_page_response

        for each_topic in json_response['content_overview']:
            # detail_adder_prompt = f'You are a professional report writer tasked with generating detailed and engaging content for a multi-page report. The report is about the online presence and activities of a specific target over the last 7 days. You will be provided with a heading, subheadings, and a brief summary. Your job is to expand this information into a comprehensive, well-structured, and detailed section. Details: - Heading: {each_topic['heading']} - Subheadings: {each_topic['subheadings']} - Summary: {each_topic['summary']}Instructions: 1. Write a detailed introduction for the section based on the heading and summary. 2. Expand each subheading into a detailed subsection: - Include numerical data, insights, and analysis where applicable. - Provide comparisons, psychometric insights, and examples. - Explain the significance of the subtopic and its impact. 3. Conclude the section with a summary of key takeaways and implications. Special Instructions: - Ensure the tone is professional and engaging. - Use data-driven insights and concrete examples to support your writing. - Write content that flows naturally and connects ideas between subsections.'

            detail_adder_prompt = f'You are a professional report writer tasked with producing highly engaging, data-driven, and contextually rich sections for a report. Each topic should feel tailored and insightful, with the details woven naturally into the narrative to maintain relevance and interest. The content must be easy to analyze while keeping the reader engaged. Guidelines for Writing: Contextual Introduction: Start with a brief but engaging introduction that sets the context for the topic and its significance. Data Integration: Incorporate quantitative data (e.g., percentages, growth rates, and engagement metrics) into the narrative. Use platform-specific insights where relevant, but only if supported by the data. Provide comparative metrics when logical (e.g., previous periods, competitors, or industry benchmarks). Dynamic Structure: Let the structure adapt to the topic’s unique aspects. For example: If sentiment analysis is central, focus on emotional dynamics and trends. If platform-specific activity is significant, emphasize platform breakdowns. If actionable recommendations are clear, make them the highlight. Avoid forcing any data or subsections if the input doesn’t support it; instead, highlight what is significant or missing. Style: Blend narrative and analysis to keep the content informative and engaging. Use professional but conversational language where appropriate. Focus on creating an immersive experience for the reader. Task: Based on the given {each_topic['heading']}, {each_topic['subheadings']}, and {each_topic['summary']}, expand the information into a well-rounded, data-rich section. Make the content dynamic, logical, and engaging, emphasizing actionable insights and measurable trends wherever they naturally fit. Avoid unnecessary rigidity; prioritize meaningful details that make the report compelling.'

            output = chat_assist(detail_adder_instructions, detail_adder_prompt)
            full_report += output
        

        summary_prompt = "Summarize the provided text into a concise, maximum 2-page report. Include key insights organized into bullet points, ensure quantitative data is highlighted, and cover major topics such as public sentiment, humanitarian impact, community reactions, and key challenges with recommendations for action. If the text references any steps the U.S. government is taking to address the challenges, include those details; otherwise, omit any mention of potential U.S. government actions or solutions."

        summary = chat_assist("None", full_report + summary_prompt)

    except Exception as e:
        app.logger.error(f"Error occured: {e}")

    

    return save_to_doc(summary, topic_name, region)


    # if request.method == 'POST':
    #     # Retrieve user input from the floating text box
    #     user_input = request.form.get('user_input', '')
    #     # Generate a report (or process the input)
    #     html_report = markdown.markdown(f"Processed input: {user_input}")  # Replace with actual processing logic
    # else:
    #     html_report = "No input provided yet. Submit text to see the report."
    #     file = open("output.html","r")
    #     html_report = file.read()

def save_to_doc(full_report, topic_name, region):

    html_report = markdown.markdown(full_report)
    soup = BeautifulSoup(html_report, 'html.parser')
    doc = Document()

    for element in soup.contents:
        if element.name == 'h1':
            heading = doc.add_heading(level=1)
            run = heading.add_run(element.text)
            run.font.color.rgb = RGBColor(1, 1, 60)
            run.font.size = Pt(24)
        elif element.name == 'h2':
            heading = doc.add_heading(level=2)
            run = heading.add_run(element.text)
            run.font.color.rgb = RGBColor(1, 1, 60)
            run.font.size = Pt(20)
        elif element.name == 'h3':
            heading = doc.add_heading(level=3)
            run = heading.add_run(element.text)
            run.font.color.rgb = RGBColor(1, 1, 60)
            run.font.size = Pt(16)
        elif element.name == 'h4':
            heading = doc.add_heading(level=4)
            run = heading.add_run(element.text)
            run.font.color.rgb = RGBColor(1, 1, 60)
            run.font.size = Pt(14)
        elif element.name == 'p':
            heading = doc.add_paragraph()
            run = heading.add_run(element.text)
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.size = Pt(11)
        elif element.name == 'ul':
            for li in element.find_all('li'):
                doc.add_paragraph(li.text, style='List Bullet')
    # Save the document to a BytesIO stream
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    app.logger.info("Sending response file")
    return send_file(
        file_stream,
        as_attachment=True,
        download_name=f"{topic_name}_{region}_report.docx",
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

def chat_assist(instructions, prompt):
    chat_response = client.chat.completions.create(
        model = "gpt-4o",
        messages = [
            {"role": "system", "content": instructions },
            {"role": "user", "content": prompt }
        ]
    )
    return chat_response.choices[0].message.content


def google_results(query):
    api_key = os.getenv("GOOGLE_SEARCH_API")
    search_engine_id = os.getenv("SEARCH_ENGINE_ID")
    service = build("customsearch", "v1", developerKey=api_key)
    
    try:
        result = service.cse().list(
            q=query,  # Search query
            cx=search_engine_id,     # Custom Search Engine ID
            dateRestrict = "w1",
        ).execute()
        # Extract the results
        if "items" in result:
            print([item["link"] for item in result["items"]])
            return [item["link"] for item in result["items"]]
        else:
            print("No results found in the response.")
            return None

    except Exception as e:
        print("An error occurred:", str(e))
        return None


@app.route('/generate', methods=['GET', 'POST'])
def generate():
    if request.method == 'POST':
        # Simulate a long-running process
        time.sleep(5)  # Replace this with the actual long process
        return jsonify({"status": "completed", "message": "Document is ready!"})
    return render_template('loading.html')

if __name__ == '__main__':
    app.run(debug=True)